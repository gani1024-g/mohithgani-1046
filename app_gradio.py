"""DocuSense AI - Document Review & Analysis Workspace.
PostgreSQL + pgvector backed application.
"""

import base64
import html
import io
import json
import os
import re
from datetime import datetime, timedelta
import threading
import time
from typing import Any, Dict, Generator, List, Optional, Tuple

import bcrypt
import pymupdf
import gradio as gr
import numpy as np
from dotenv import load_dotenv
from groq import Groq
from PIL import Image
import pytesseract
from sentence_transformers import SentenceTransformer
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from models import (
    Base,
    ChatMessage,
    ChatSession,
    Document,
    DocumentChunk,
    DocumentSummary,
    MessageCitation,
    User,
)
from ingestion_pipeline import process_document_pipeline, validate_file_magic

load_dotenv()

# =====================================================================
# 1. Configuration
# =====================================================================
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise RuntimeError("GROQ_API_KEY environment variable is not set. Add it to .env.")

groq_client = Groq(api_key=GROQ_API_KEY)
MODEL_NAME = "openai/gpt-oss-120b"

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql+psycopg://postgres:YOUR_POSTGRES_PASSWORD@127.0.0.1:5432/grounded_rag",
)

print("Connecting to PostgreSQL + pgvector...")

engine = create_engine(
    DATABASE_URL,
    echo=False,
    pool_pre_ping=True,
)

SessionLocal = sessionmaker(
    bind=engine,
    expire_on_commit=False,
)

with engine.begin() as conn:
    conn.exec_driver_sql("CREATE EXTENSION IF NOT EXISTS vector")
    # Safe, idempotent migration for temporary guest-account expiry.
    conn.exec_driver_sql('ALTER TABLE users ADD COLUMN IF NOT EXISTS "ExpiresAt" TIMESTAMP NULL')
    conn.exec_driver_sql('CREATE INDEX IF NOT EXISTS ix_users_ExpiresAt ON users ("ExpiresAt")')

Base.metadata.create_all(engine)
print("PostgreSQL connected and tables verified.")

print("Loading BGE-large-en Embedding Model...")
embedding_model = SentenceTransformer("BAAI/bge-large-en-v1.5")

MAX_FILE_SIZE_MB = 50
MAX_PAGE_COUNT = 300
MAX_GUEST_USES = 5
GUEST_RETENTION_HOURS = 2
GUEST_CLEANUP_INTERVAL_SECONDS = 15 * 60
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx", ".jpg", ".jpeg", ".png"}
MIME_TYPE_MAPPING = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}

# =====================================================================
# 2. Guest lifecycle / cleanup
# =====================================================================
def cleanup_expired_guest_data() -> int:
    """Delete expired temporary guest users and all data owned by them."""
    now = datetime.utcnow()
    deleted = 0
    try:
        with SessionLocal() as db:
            guests = db.execute(
                select(User).where(
                    User.auth_provider == "Guest",
                    User.expires_at.is_not(None),
                    User.expires_at < now,
                )
            ).scalars().all()

            for guest in guests:
                db.delete(guest)
                deleted += 1

            if deleted:
                db.commit()
                print(f"Cleaned up {deleted} expired guest account(s).")
            return deleted
    except Exception as exc:
        print(f"Guest cleanup error: {exc}")
        return 0


def touch_guest_activity(user_data: Optional[Dict]) -> bool:
    """Extend a guest session by GUEST_RETENTION_HOURS after activity."""
    if not user_data or str(user_data.get("role", "")).upper() != "GUEST":
        return True

    try:
        user_id = int(user_data["id"])
        now = datetime.utcnow()
        with SessionLocal() as db:
            guest = db.execute(select(User).where(User.id == user_id)).scalar_one_or_none()
            if not guest or guest.auth_provider != "Guest" or not guest.is_active:
                return False
            if guest.expires_at and guest.expires_at < now:
                return False
            guest.expires_at = now + timedelta(hours=GUEST_RETENTION_HOURS)
            db.commit()
            return True
    except Exception as exc:
        print(f"Guest activity update error: {exc}")
        return False


def guest_cleanup_worker() -> None:
    """Background cleanup so expired guest data is removed without a restart."""
    while True:
        try:
            cleanup_expired_guest_data()
        except Exception as exc:
            print(f"Guest cleanup worker error: {exc}")
        time.sleep(GUEST_CLEANUP_INTERVAL_SECONDS)


# Clean old guest data once at startup, then continue every 15 minutes.
cleanup_expired_guest_data()
threading.Thread(target=guest_cleanup_worker, daemon=True, name="guest-cleanup").start()


# =====================================================================
# 3. Authentication
# =====================================================================
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def verify_password(plain_password: str, hashed_password: str) -> bool:
    if not hashed_password:
        return False
    return bcrypt.checkpw(plain_password.encode("utf-8"), hashed_password.encode("utf-8"))


def auth_register(email: str, full_name: str, password: str, confirm_password: str, role: str):
    clean_email = (email or "").strip().lower()
    clean_name = (full_name or "").strip()
    if not clean_email or not clean_name or not password or not confirm_password:
        return "Please fill in all required fields.", None
    if password != confirm_password:
        return "Passwords do not match. Please re-enter.", None
    if len(password) < 6:
        return "Password must be at least 6 characters long.", None

    try:
        with SessionLocal() as db:
            existing = db.execute(select(User).where(User.email == clean_email)).scalar_one_or_none()
            if existing:
                return f"An account with email '{clean_email}' already exists.", None

            user = User(
                email=clean_email,
                full_name=clean_name,
                role=(role or "USER").upper(),
                hashed_password=hash_password(password),
                auth_provider="Local",
                is_active=True,
            )
            db.add(user)
            db.commit()
            db.refresh(user)
            user_data = {"id": user.id, "email": user.email, "full_name": user.full_name, "role": user.role}
            return "Account created successfully! Logging you in...", user_data
    except Exception as exc:
        return f"Database Error: {exc}", None


def auth_login(email: str, password: str):
    clean_email = (email or "").strip().lower()
    if not clean_email or not password:
        return "Please enter both email and password.", None
    try:
        with SessionLocal() as db:
            user = db.execute(select(User).where(User.email == clean_email)).scalar_one_or_none()
            if not user or not user.is_active:
                return "No active account found with this email.", None
            if not verify_password(password, user.hashed_password or ""):
                return "Incorrect password.", None
            user.last_login_at = datetime.utcnow()
            db.commit()
            data = {"id": user.id, "email": user.email, "full_name": user.full_name, "role": user.role}
            return f"Welcome back, {user.full_name}!", data
    except Exception as exc:
        return f"Database Error: {exc}", None


def auth_guest_login():
    guest_email = f"guest_{os.urandom(8).hex()}@docusense.local"
    guest_name = f"Guest-{os.urandom(3).hex().upper()}"
    try:
        with SessionLocal() as db:
            guest = User(
                email=guest_email,
                full_name=guest_name,
                role="GUEST",
                hashed_password=hash_password(os.urandom(16).hex()),
                auth_provider="Guest",
                is_active=True,
                expires_at=datetime.utcnow() + timedelta(hours=GUEST_RETENTION_HOURS),
            )
            db.add(guest)
            db.commit()
            db.refresh(guest)
            data = {"id": guest.id, "email": guest.email, "full_name": guest.full_name, "role": "GUEST"}
            return "Entered workspace as Guest!", data
    except Exception as exc:
        return f"Database Error: {exc}", None


def get_user_documents(user_id: int):
    if not user_id:
        return []
    try:
        with SessionLocal() as db:
            docs = db.execute(
                select(Document)
                .where(Document.user_id == int(user_id))
                .order_by(Document.uploaded_at.desc())
            ).scalars().all()
            return [
                (f"{d.title} ({d.page_count} page{'s' if d.page_count != 1 else ''})", str(d.id))
                for d in docs
            ]
    except Exception as exc:
        print(f"Error fetching user documents: {exc}")
        return []


# =====================================================================
# 3. Guest usage
# =====================================================================
def consume_guest_use(user_data: Optional[Dict], usage_count: int) -> Tuple[int, bool]:
    current = int(usage_count or 0)
    if user_data and user_data.get("role") == "GUEST":
        if current >= MAX_GUEST_USES:
            return current, False
        return current + 1, True
    return current, True


def guest_limit_message() -> str:
    return "Guest access limit reached (5 actions). Please create a free account to continue."


# =====================================================================
# 4. Non-PDF extraction helpers
# =====================================================================
def extract_non_pdf_chunks(file_bytes: bytes, filename: str, ext: str, doc_id: int) -> Tuple[int, List[Dict]]:
    chunks: List[Dict] = []

    if ext == ".txt":
        text_content = file_bytes.decode("utf-8", errors="ignore")
        words = text_content.split()
        if not words:
            return 1, []
        for idx, start in enumerate(range(0, len(words), 768)):
            content = " ".join(words[start:start + 768]).strip()
            if content:
                chunks.append(
                    {
                        "doc_id": doc_id,
                        "chunk_index": idx,
                        "page_num": 1 + start // 768,
                        "content": content,
                        "bbox_json": json.dumps({"x0": 0, "y0": 0, "x1": 612, "y1": 792}),
                        "token_count": len(content.split()),
                    }
                )
        return max(1, len(chunks)), chunks

    if ext == ".docx":
        try:
            import docx
        except ImportError as exc:
            raise ValueError("python-docx is not installed. Install it with: pip install python-docx") from exc
        document = docx.Document(io.BytesIO(file_bytes))
        text_content = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
        words = text_content.split()
        if not words:
            return 1, []
        for idx, start in enumerate(range(0, len(words), 768)):
            content = " ".join(words[start:start + 768]).strip()
            chunks.append(
                {
                    "doc_id": doc_id,
                    "chunk_index": idx,
                    "page_num": idx + 1,
                    "content": content,
                    "bbox_json": json.dumps({"x0": 0, "y0": 0, "x1": 612, "y1": 792}),
                    "token_count": len(content.split()),
                }
            )
        return len(chunks), chunks

    if ext in {".png", ".jpg", ".jpeg"}:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        ocr_text = pytesseract.image_to_string(image).strip()
        if not ocr_text:
            ocr_text = f"Visual document image: {filename}"
        chunks.append(
            {
                "doc_id": doc_id,
                "chunk_index": 0,
                "page_num": 1,
                "content": ocr_text,
                "bbox_json": json.dumps({"x0": 0, "y0": 0, "x1": image.width, "y1": image.height}),
                "token_count": len(ocr_text.split()),
            }
        )
        return 1, chunks

    raise ValueError(f"Unsupported file format '{ext}'.")


def add_embeddings(chunks: List[Dict]) -> None:
    if not chunks:
        return
    texts = [c["content"] for c in chunks]
    embeddings = embedding_model.encode(
        texts,
        normalize_embeddings=True,
        show_progress_bar=False,
    )
    for chunk, emb in zip(chunks, embeddings):
        chunk["embedding"] = [float(value) for value in emb]



# =====================================================================
# 5. Upload + database storage
# =====================================================================
def normalize_gradio_file(file_obj: Any) -> Tuple[str, str]:
    if isinstance(file_obj, (tuple, list)):
        file_obj = next((x for x in file_obj if x), None)
    if isinstance(file_obj, dict):
        path = file_obj.get("path") or file_obj.get("filepath") or file_obj.get("name")
        name = file_obj.get("orig_name") or file_obj.get("original_name") or file_obj.get("name")
    else:
        path = getattr(file_obj, "path", None) or getattr(file_obj, "filepath", None) or getattr(file_obj, "name", None) or file_obj
        name = getattr(file_obj, "orig_name", None) or getattr(file_obj, "original_name", None)

    if not path:
        raise ValueError("File path could not be determined.")
    path = str(os.fspath(path))
    name = os.path.basename(str(name or path))
    return path, name


def process_uploaded_document(file_obj: Any, user_data: Optional[Dict]):
    if not user_data or not user_data.get("id"):
        return "Please sign in or continue as guest first.", None, None, gr.update()
    if not file_obj:
        return "No file selected.", None, None, gr.update()
    if not touch_guest_activity(user_data):
        return "Guest session expired. Please continue as guest again.", None, None, gr.update()

    path, filename = normalize_gradio_file(file_obj)
    if not os.path.isfile(path):
        raise ValueError(f"File could not be accessed at: {path}")

    with open(path, "rb") as handle:
        file_bytes = handle.read()

    file_size_bytes = len(file_bytes)
    if file_size_bytes > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise ValueError(f"File size exceeds maximum allowed limit of {MAX_FILE_SIZE_MB} MB.")

    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file format '{ext}'.")

    mime_type = MIME_TYPE_MAPPING[ext]
    if ext == ".pdf":
        # Strong magic-byte validation before anything is inserted.
        actual_mime = validate_file_magic(file_bytes, MAX_FILE_SIZE_MB)
        if actual_mime != "application/pdf":
            raise ValueError(f"The uploaded file is not a valid PDF (detected {actual_mime}).")

    with SessionLocal() as db:
        # Create the document first so PostgreSQL generates DocumentID.
        doc = Document(
            user_id=int(user_data["id"]),
            title=filename,
            original_name=filename,
            file_size_bytes=file_size_bytes,
            page_count=1,
            chunk_count=0,
            storage_path=None,
            mime_type=mime_type,
            status="Processing",
            file_data=file_bytes,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        doc_id = doc.id

        try:
            if ext == ".pdf":
                page_count, chunks = process_document_pipeline(
                    doc_id=doc_id,
                    file_bytes=file_bytes,
                    embedding_model=embedding_model,
                )
            else:
                page_count, chunks = extract_non_pdf_chunks(file_bytes, filename, ext, doc_id)
                add_embeddings(chunks)

            if page_count > MAX_PAGE_COUNT:
                raise ValueError(f"Document has {page_count} pages (limit is {MAX_PAGE_COUNT}).")
            if not chunks:
                raise ValueError("No readable text was found in the uploaded document.")

            for chunk in chunks:
                db.add(
                    DocumentChunk(
                        doc_id=doc_id,
                        chunk_index=int(chunk.get("chunk_index", 0)),
                        page_num=int(chunk.get("page_num", 1)),
                        content=chunk["content"],
                        token_count=chunk.get("token_count"),
                        embedding_id="BAAI/bge-large-en-v1.5",
                        embedding=chunk["embedding"],
                        bbox_json=chunk.get("bbox_json"),
                    )
                )

            doc.page_count = page_count
            doc.chunk_count = len(chunks)
            doc.status = "Indexed"
            doc.updated_at = datetime.utcnow()
            db.commit()

            status_msg = f"Indexed **{filename}** ({page_count} page{'s' if page_count != 1 else ''}, {len(chunks)} chunks)."
            docs = get_user_documents(int(user_data["id"]))
            return status_msg, str(doc_id), str(doc_id), gr.update(choices=docs, value=str(doc_id))

        except Exception:
            doc.status = "Failed"
            doc.updated_at = datetime.utcnow()
            db.commit()
            raise


# =====================================================================
# 6. File preview directly from PostgreSQL BYTEA FileData
# =====================================================================
def render_file_bytes(file_bytes: bytes, mime_type: str, filename: str) -> str:
    if not file_bytes:
        return "<div class='pdf-frame-box empty-preview'>No stored file data.</div>"

    if mime_type == "application/pdf":
        encoded = base64.b64encode(file_bytes).decode("ascii")
        return (
            "<div class='pdf-frame-box' style='height:480px;border-radius:8px;overflow:hidden;'>"
            f"<iframe src='data:application/pdf;base64,{encoded}' width='100%' height='480px' "
            "style='border:none;width:100%;height:100%;' type='application/pdf'></iframe>"
            "</div>"
        )

    if mime_type.startswith("image/"):
        encoded = base64.b64encode(file_bytes).decode("ascii")
        return (
            "<div class='pdf-frame-box' style='height:480px;display:flex;align-items:center;justify-content:center;background:#080c14;padding:12px;'>"
            f"<img src='data:{mime_type};base64,{encoded}' style='max-width:100%;max-height:100%;object-fit:contain;border-radius:6px;' />"
            "</div>"
        )

    if mime_type == "text/plain":
        text_content = html.escape(file_bytes.decode("utf-8", errors="ignore")[:15000])
        return (
            "<div class='pdf-frame-box' style='height:480px;overflow-y:auto;padding:18px;background:#0c1220;"
            "color:#cbd5e1;font-family:monospace;font-size:13px;line-height:1.5;white-space:pre-wrap;'>"
            f"{text_content}</div>"
        )

    return (
        "<div class='pdf-frame-box' style='height:480px;display:flex;align-items:center;justify-content:center;"
        "color:#94a3b8;'>Stored document: " + html.escape(filename) + "</div>"
    )


def get_document_for_user(doc_id: int, user_data: Dict) -> Optional[Document]:
    with SessionLocal() as db:
        doc = db.execute(select(Document).where(Document.id == int(doc_id))).scalar_one_or_none()
        if not doc:
            return None
        role = str(user_data.get("role", "")).upper()
        if role == "GUEST":
            if not touch_guest_activity(user_data):
                return None
            if int(doc.user_id) != int(user_data.get("id")):
                return None
        elif int(doc.user_id) != int(user_data.get("id")) and role != "ADMIN":
            return None
        return doc


def preview_document_from_db(doc_id: Optional[str], user_data: Optional[Dict]) -> str:
    if not doc_id or not user_data:
        return "<div class='pdf-frame-box empty-preview'>No document selected</div>"
    try:
        with SessionLocal() as db:
            doc = db.execute(select(Document).where(Document.id == int(doc_id))).scalar_one_or_none()
            if not doc:
                return "<div class='pdf-frame-box empty-preview'>Document not found</div>"
            role = str(user_data.get("role", "")).upper()
            if role == "GUEST":
                if not touch_guest_activity(user_data) or int(doc.user_id) != int(user_data.get("id")):
                    return "<div class='pdf-frame-box empty-preview'>Guest session expired or access denied</div>"
            elif int(doc.user_id) != int(user_data.get("id")) and role != "ADMIN":
                return "<div class='pdf-frame-box empty-preview'>Access denied</div>"
            return render_file_bytes(doc.file_data or b"", doc.mime_type, doc.original_name)
    except Exception as exc:
        return f"<div class='pdf-frame-box error-preview'>Error loading preview: {html.escape(str(exc))}</div>"


# =====================================================================
# 7. RAG helpers
# =====================================================================
def allowed_document(db, doc_id: int, user_data: Dict) -> Optional[Document]:
    doc = db.execute(select(Document).where(Document.id == int(doc_id))).scalar_one_or_none()
    if not doc:
        return None
    role = str(user_data.get("role", "")).upper()
    if role == "GUEST":
        if not touch_guest_activity(user_data):
            return None
        if int(doc.user_id) != int(user_data.get("id")):
            return None
    else:
        is_admin = role == "ADMIN"
        if int(doc.user_id) != int(user_data.get("id")) and not is_admin:
            return None
    return doc


def get_or_create_chat_session(db, doc_id: int, user_id: int) -> ChatSession:
    session = db.execute(
        select(ChatSession)
        .where(ChatSession.doc_id == int(doc_id), ChatSession.user_id == int(user_id))
        .order_by(ChatSession.updated_at.desc())
    ).scalars().first()
    if session:
        return session
    session = ChatSession(doc_id=int(doc_id), user_id=int(user_id), title="Document Conversation")
    db.add(session)
    db.flush()
    return session


def stream_rag_response(
    user_query: str,
    chat_history: list,
    doc_id_str: str,
    user_data: Optional[Dict],
    usage_count: int,
) -> Generator[Tuple[list, str, int], None, None]:
    if not (user_query or "").strip():
        yield chat_history or [], "", usage_count
        return
    chat_history = list(chat_history or [])
    if not user_data:
        chat_history += [{"role": "user", "content": user_query}, {"role": "assistant", "content": "Please sign in or continue as guest first."}]
        yield chat_history, "", usage_count
        return
    if not doc_id_str:
        chat_history += [{"role": "user", "content": user_query}, {"role": "assistant", "content": "Please select or upload a document first."}]
        yield chat_history, "", usage_count
        return

    usage_count, allowed = consume_guest_use(user_data, usage_count)
    if not allowed:
        chat_history += [{"role": "user", "content": user_query}, {"role": "assistant", "content": guest_limit_message()}]
        yield chat_history, "", usage_count
        return

    try:
        with SessionLocal() as db:
            doc = allowed_document(db, int(doc_id_str), user_data)
            if not doc:
                chat_history += [{"role": "user", "content": user_query}, {"role": "assistant", "content": "Access Denied: You do not own this document."}]
                yield chat_history, "", usage_count
                return

            q_vec = np.asarray(
                embedding_model.encode(
                    user_query,
                    normalize_embeddings=True,
                    show_progress_bar=False,
                ),
                dtype=np.float32,
            ).tolist()

            top_chunks = db.execute(
                select(DocumentChunk)
                .where(DocumentChunk.doc_id == doc.id)
                .order_by(DocumentChunk.embedding.cosine_distance(q_vec))
                .limit(4)
            ).scalars().all()

        if not top_chunks:
            chat_history += [
                {"role": "user", "content": user_query},
                {"role": "assistant", "content": "No indexed content chunks found for this document."},
            ]
            yield chat_history, "", usage_count
            return

        context = "\n\n".join(
            f"[Page {c.page_num}]\n{c.content}" for c in top_chunks
        )

        system_prompt = (
            "You are DocuSense AI, an intelligent document review assistant.\n"
            "Answer strictly from the provided document context. Cite page numbers as [Page X].\n"
            "If the context does not contain the answer, say so clearly.\n\n"
            f"Document context:\n{context}"
        )

        chat_history.append({"role": "user", "content": user_query})
        chat_history.append({"role": "assistant", "content": ""})
        yield chat_history, "", usage_count

        completion = groq_client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_query}],
            temperature=0.1,
            max_tokens=1024,
            stream=True,
        )
        answer = ""
        for part in completion:
            delta = part.choices[0].delta.content or ""
            answer += delta
            chat_history[-1]["content"] = answer
            yield chat_history, "", usage_count

        with SessionLocal() as db:
            session = get_or_create_chat_session(
                db, int(doc_id_str), int(user_data["id"])
            )

            user_msg = ChatMessage(
                session_id=session.session_id,
                role="user",
                content=user_query,
            )
            db.add(user_msg)
            db.flush()

            assistant_msg = ChatMessage(
                session_id=session.session_id,
                role="assistant",
                content=answer,
            )
            db.add(assistant_msg)
            db.flush()

            for citation_order, chunk in enumerate(top_chunks):
                db.add(
                    MessageCitation(
                        message_id=assistant_msg.message_id,
                        chunk_id=chunk.chunk_id,
                        citation_order=citation_order,
                    )
                )

            session.updated_at = datetime.utcnow()
            db.commit()

    except Exception as exc:
        chat_history.append({"role": "assistant", "content": f"Error generating response: {exc}"})
        yield chat_history, "", usage_count


# =====================================================================
# 8. Summaries
# =====================================================================
def parse_llm_json(text_str: str) -> Dict:
    if not text_str:
        return {}
    cleaned = re.sub(r"```(?:json)?", "", text_str).strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        try:
            return json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError:
            pass
    return {}


def generate_summary_and_risks(
    doc_id_str: str,
    user_data: Optional[Dict],
    usage_count: int,
):
    """Generate a detailed executive summary, multiple key takeaways, and risks."""
    if not user_data:
        return "*Please sign in or continue as guest first.*", "", "", usage_count

    if not doc_id_str:
        return "*Please select a document first.*", "", "", usage_count

    usage_count, allowed = consume_guest_use(user_data, usage_count)
    if not allowed:
        return guest_limit_message(), "", "", usage_count

    try:
        doc_id = int(doc_id_str)

        with SessionLocal() as db:
            doc = allowed_document(db, doc_id, user_data)
            if not doc:
                return "*Access Denied: You do not own this document.*", "", "", usage_count

            chunks = db.execute(
                select(DocumentChunk)
                .where(DocumentChunk.doc_id == doc.id)
                .order_by(DocumentChunk.chunk_index.asc())
            ).scalars().all()

        if not chunks:
            return "No text chunks found in database.", "", "", usage_count

        # Use as much indexed document text as practical for a richer summary.
        full_text = "\n\n".join(
            c.content.strip() for c in chunks if c.content and c.content.strip()
        )
        full_text = full_text[:18000]

        if not full_text.strip():
            return "No readable text found in the document.", "", "", usage_count

        summary_prompt = f"""
Create a detailed, document-grounded review of the source below.

Return ONLY valid JSON with exactly this structure:

{{
  "executive_summary": "A detailed 500-800 word synthesis covering the document's purpose, major sections, important facts, notable skills/requirements/commitments, and overall meaning.",
  "key_takeaways": [
    "Takeaway 1",
    "Takeaway 2",
    "Takeaway 3",
    "Takeaway 4",
    "Takeaway 5",
    "Takeaway 6",
    "Takeaway 7",
    "Takeaway 8"
  ],
  "risks_and_liabilities": [
    {{
      "risk_level": "HIGH",
      "clause": "Short risk title",
      "description": "Specific explanation grounded in the document"
    }}
  ]
}}

Rules:
- Use ONLY facts and statements supported by the document text.
- Do not invent missing details, names, dates, numbers, qualifications, or obligations.
- executive_summary should be detailed and organized as a coherent synthesis, not a one-sentence abstract.
- Aim for 500-800 words when the document contains enough information; for a short document, be as detailed as the source permits without repeating filler.
- key_takeaways should contain 6-8 distinct, useful points. Do not repeat the same idea in different wording.
- Prioritize concrete facts, sections, achievements, requirements, decisions, deliverables, skills, or obligations.
- risks_and_liabilities should contain 0-6 items. Include a risk only when the document supports it. Do not manufacture risks just to fill the array.
- risk_level must be exactly HIGH, MED, or LOW.
- Every risk item must be an object with risk_level, clause, and description.
- Return plain strings in key_takeaways.
- Do not return markdown.
- Do not return code fences.
- Return only the JSON object.

Document:
{full_text}
"""

        response = groq_client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {
                    "role": "system",
                    "content": "You are a meticulous document analyst. Return only valid JSON and never add unsupported facts.",
                },
                {"role": "user", "content": summary_prompt},
            ],
            temperature=0.15,
            max_tokens=1800,
        )

        raw_response = response.choices[0].message.content or ""
        data = parse_llm_json(raw_response)
        if not isinstance(data, dict):
            data = {}

        exec_summary = str(
            data.get("executive_summary", "No executive summary generated.")
        ).strip()

        raw_takeaways = data.get("key_takeaways", [])
        takeaways_list = raw_takeaways if isinstance(raw_takeaways, list) else [raw_takeaways]
        takeaways_list = [
            str(item).strip()
            for item in takeaways_list
            if str(item).strip()
        ]

        # Keep the UI focused: at most 8 distinct takeaways.
        takeaways_list = takeaways_list[:8]

        raw_risks = data.get("risks_and_liabilities", [])
        risk_items = raw_risks if isinstance(raw_risks, list) else [raw_risks]

        takeaways = "".join(
            f"<div class='takeaway-card'>"
            f"<div class='takeaway-title'>Key Point {i + 1}</div>"
            f"{html.escape(item)}"
            f"</div>"
            for i, item in enumerate(takeaways_list)
        ) or "<div class='takeaway-card'>No key takeaways identified.</div>"

        risks_html: List[str] = []

        for item in risk_items[:6]:
            # The LLM may occasionally return a string instead of an object.
            if isinstance(item, dict):
                level = str(item.get("risk_level", "MED")).upper().strip()
                clause = str(item.get("clause", "Observation")).strip()
                desc = str(item.get("description", "")).strip()
            else:
                level = "MED"
                clause = "Observation"
                desc = str(item).strip()

            if level not in {"HIGH", "MED", "LOW"}:
                level = "MED"

            badge_class = (
                "badge-high"
                if level == "HIGH"
                else "badge-med"
                if level == "MED"
                else "badge-low"
            )

            risks_html.append(
                f"<div class='risk-row'>"
                f"<span class='{badge_class}'>{html.escape(level)}</span> "
                f"<strong>{html.escape(clause)}</strong>"
                f"<div>{html.escape(desc)}</div>"
                f"</div>"
            )

        risks = "".join(risks_html) or (
            "<div class='risk-row'>No critical liabilities identified in the document.</div>"
        )

        # Persist the generated summary in DocumentSummaries.
        with SessionLocal() as db:
            existing = db.execute(
                select(DocumentSummary)
                .where(
                    DocumentSummary.doc_id == doc_id,
                    DocumentSummary.summary_mode == "Executive",
                )
            ).scalars().first()

            payload = json.dumps(data, ensure_ascii=False)

            if existing:
                existing.summary_text = exec_summary
                existing.key_takeaways_json = json.dumps(
                    takeaways_list, ensure_ascii=False
                )
                existing.risks_json = json.dumps(
                    risk_items[:6], ensure_ascii=False
                )
                existing.multimode_data_json = payload
                existing.updated_at = datetime.utcnow()
            else:
                db.add(
                    DocumentSummary(
                        doc_id=doc_id,
                        summary_mode="Executive",
                        summary_text=exec_summary,
                        key_takeaways_json=json.dumps(
                            takeaways_list, ensure_ascii=False
                        ),
                        risks_json=json.dumps(
                            risk_items[:6], ensure_ascii=False
                        ),
                        multimode_data_json=payload,
                    )
                )

            db.commit()

        return exec_summary, takeaways, risks, usage_count

    except Exception as exc:
        return f"Error: {exc}", "", "", usage_count


# =====================================================================
# 9. UI
# =====================================================================
custom_css = """
:root { --primary-color:#6366f1; --primary-hover:#4f46e5; --bg-dark:#090d16; --card-bg:#131b2e; --card-border:#1e293b; --text-main:#f8fafc; --text-muted:#94a3b8; }
body,.gradio-container { background-color:var(--bg-dark)!important; color:var(--text-main)!important; max-width:1440px!important; margin:0 auto!important; }
footer,.footer,.gradio-footer,[data-testid="footer"],.show-api,.api-docs { display:none!important; visibility:hidden!important; height:0!important; }
.doc-header-box { display:flex; justify-content:space-between; align-items:center; padding:18px 24px; background:linear-gradient(135deg,rgba(30,41,59,.7),rgba(15,23,42,.85))!important; border:1px solid rgba(255,255,255,.08)!important; border-radius:16px!important; margin-bottom:20px!important; }
.doc-title { font-size:24px!important; font-weight:700!important; color:#f8fafc!important; }
.user-badge-box { display:flex; flex-direction:column; align-items:flex-end; justify-content:center; gap:6px; }
.takeaway-card { background:linear-gradient(135deg,#131b2e,#17223b); border:1px solid #1e293b; border-left:4px solid #6366f1; border-radius:10px; padding:14px 18px; margin-bottom:12px; color:#e2e8f0; font-size:14px; line-height:1.6; }
.takeaway-title { font-weight:700; color:#a5b4fc; margin-bottom:4px; font-size:13px; text-transform:uppercase; }
.risk-row { background:#131b2e; border:1px solid #1e293b; border-radius:10px; padding:14px 18px; margin-bottom:12px; }
.badge-high,.badge-med,.badge-low { padding:3px 9px; border-radius:6px; font-size:11px; font-weight:700; display:inline-block; margin-bottom:4px; }
.badge-high { background:rgba(239,68,68,.15); color:#fca5a5; border:1px solid rgba(239,68,68,.4); }
.badge-med { background:rgba(245,158,11,.15); color:#fde68a; border:1px solid rgba(245,158,11,.4); }
.badge-low { background:rgba(16,185,129,.15); color:#a7f3d0; border:1px solid rgba(16,185,129,.4); }
.pdf-frame-box { background:#0b0f19; border:1px solid #1e293b; border-radius:12px; overflow:hidden; }
.empty-preview { height:480px; display:flex; align-items:center; justify-content:center; color:#94a3b8; }
.error-preview { height:480px; display:flex; align-items:center; justify-content:center; color:#ef4444; padding:20px; }
"""

with gr.Blocks(
    title="DocuSense AI",
    css=custom_css,
    head="<title>DocuSense AI</title>"
) as demo:
    current_user_state = gr.State(value=None)
    active_doc_id = gr.State(value=None)
    guest_usage_state = gr.State(value=0)

    with gr.Row(elem_classes=["doc-header-box"]):
        with gr.Column(scale=8, min_width=200):
            gr.HTML("<div class='doc-title'>DocuSense AI — Document Review &amp; Analysis</div>")
        with gr.Column(scale=4, min_width=200, elem_classes=["user-badge-box"]):
            user_badge = gr.Markdown("*Not Signed In*", elem_id="user_badge")
            btn_logout = gr.Button("Sign Out", visible=False, variant="secondary", size="sm")

    with gr.Group(visible=True) as auth_view:
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("### Access Workspace")
                with gr.Tabs():
                    with gr.TabItem("Sign In"):
                        login_email = gr.Textbox(label="Email", placeholder="name@company.com")
                        login_password = gr.Textbox(label="Password", type="password")
                        show_login_pass = gr.Checkbox(label="Show Password", value=False)
                        btn_do_login = gr.Button("Sign In", variant="primary")
                        login_status = gr.Markdown()
                    with gr.TabItem("Create Account"):
                        reg_name = gr.Textbox(label="Full Name", placeholder="Jane Doe")
                        reg_email = gr.Textbox(label="Email", placeholder="name@company.com")
                        reg_password = gr.Textbox(label="Password", type="password")
                        reg_confirm_password = gr.Textbox(label="Confirm Password", type="password")
                        show_reg_pass = gr.Checkbox(label="Show Passwords", value=False)
                        reg_role = gr.Dropdown(label="Role", choices=["USER", "ANALYST", "ADMIN"], value="USER")
                        btn_do_register = gr.Button("Create Account", variant="primary")
                        reg_status = gr.Markdown()
                gr.Markdown("---")
                btn_guest_access = gr.Button("Continue as Guest", variant="secondary")

    with gr.Group(visible=False) as workspace_view:
        with gr.Row():
            with gr.Column(scale=5):
                doc_selector = gr.Dropdown(label="Select Active Document", choices=[], interactive=True)
                file_input = gr.File(
                    label="Upload Document (PDF, TXT, DOCX, JPG, PNG — Max 50MB)",
                    file_types=[".pdf", ".txt", ".docx", ".jpg", ".jpeg", ".png"],
                    file_count="single",
                    type="filepath",
                )
                status_banner = gr.Markdown("*Ready to review documents.*")
                gr.Markdown("#### Document Preview")
                pdf_preview_html = gr.HTML("<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>")
                with gr.Accordion("Contextual Text Inspector", open=False):
                    snippet_input = gr.Textbox(label="Selected Text / Clause", lines=3)
                    with gr.Row():
                        btn_explain = gr.Button("Explain Simply", size="sm")
                        btn_sum_snip = gr.Button("Summarize Snippet", size="sm")
                        btn_risk_snip = gr.Button("Identify Risks", size="sm")
                    snippet_output = gr.Markdown()

            with gr.Column(scale=7):
                with gr.Tabs():
                    with gr.TabItem("Grounded RAG Chat"):
                        chatbot = gr.Chatbot(label="Conversation", height=420)
                        gr.Markdown("<small style='color:#94a3b8;font-weight:600;'>Quick Suggested Questions:</small>")
                        with gr.Row():
                            chip_1 = gr.Button("Key Takeaways", size="sm")
                            chip_2 = gr.Button("Risks & Prerequisites", size="sm")
                            chip_3 = gr.Button("Roadmap & Steps", size="sm")
                            chip_4 = gr.Button("Action Items", size="sm")
                        with gr.Row():
                            query_input = gr.Textbox(placeholder="Ask any question about the uploaded document...", scale=4, show_label=False)
                            btn_send = gr.Button("Send", variant="primary", scale=1)
                            btn_clear_chat = gr.Button("Clear", variant="secondary", scale=1)
                    with gr.TabItem("Executive Summary & Risks"):
                        with gr.Row():
                            btn_summarize = gr.Button("Generate Multi-Mode Summary", variant="primary", scale=3)
                            btn_export = gr.Button("Export Report (.md)", variant="secondary", scale=1)
                        export_file = gr.File(label="Download Generated Report", visible=False)
                        gr.Markdown("### Executive Summary")
                        out_exec_summary = gr.Markdown("*Click 'Generate Multi-Mode Summary' to synthesize the document.*")
                        gr.Markdown("### Key Takeaways")
                        out_takeaways = gr.HTML("<div style='color:#94a3b8;'>No summary generated yet.</div>")
                        gr.Markdown("### Risks & Prerequisites")
                        out_risks = gr.HTML("<div style='color:#94a3b8;'>No summary generated yet.</div>")

    # =================================================================
    # 10. Event handlers
    # =================================================================
    def toggle_login_pass_visibility(show):
        return gr.update(type="text" if show else "password")

    def toggle_reg_pass_visibility(show):
        new_type = "text" if show else "password"
        return gr.update(type=new_type), gr.update(type=new_type)

    show_login_pass.change(toggle_login_pass_visibility, [show_login_pass], [login_password])
    show_reg_pass.change(toggle_reg_pass_visibility, [show_reg_pass], [reg_password, reg_confirm_password])

    def handle_login_success(email, password):
        msg, user_data = auth_login(email, password)

        if user_data:
            docs = get_user_documents(user_data["id"])

            badge = (
                f"**{user_data['full_name']}** "
                f"(`{user_data['role']}`) | {user_data['email']}"
            )

        # IMPORTANT:
        # Do NOT automatically select or preview the previous document
        # when the user logs in.
        return (
            user_data,
            gr.update(visible=False),   # auth_view
            gr.update(visible=True),    # workspace_view
            badge,
            gr.update(visible=True, value="Sign Out"),
            gr.update(
                choices=docs,
                value=None               # nothing selected after login
            ),
            None,                       # active_doc_id
            "<div class='pdf-frame-box empty-preview'>"
            "Select or upload a document to preview"
            "</div>",
            msg,
            0,
        )
        return (
            None,
            gr.update(),
            gr.update(),
            "*Not Signed In*",
            gr.update(),
            gr.update(),
            None,
            "<div class='pdf-frame-box empty-preview'>"
            "Select or upload a document to preview"
            "</div>",
            msg,
            0,
         )
    def handle_register_success(email, name, password, confirm, role):
        msg, user_data = auth_register(email, name, password, confirm, role)
        if user_data:
            badge = f"**{user_data['full_name']}** (`{user_data['role']}`) | {user_data['email']}"
            return user_data, gr.update(visible=False), gr.update(visible=True), badge, gr.update(visible=True, value="Sign Out"), gr.update(choices=[], value=None), None, "<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>", msg, 0
        return None, gr.update(), gr.update(), "*Not Signed In*", gr.update(), gr.update(), None, "<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>", msg, 0

    def handle_guest_access():
        # Remove stale guest accounts before creating a new temporary session.
        cleanup_expired_guest_data()
        msg, user_data = auth_guest_login()
        if user_data:
            badge = f"**{user_data['full_name']}** (`{user_data['role']}`) | *Guest Session — expires after {GUEST_RETENTION_HOURS}h inactivity*"
            return user_data, gr.update(visible=False), gr.update(visible=True), badge, gr.update(visible=True, value="Exit Guest Mode"), gr.update(choices=[], value=None), None, "<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>", msg, 0
        return None, gr.update(), gr.update(), "*Not Signed In*", gr.update(), gr.update(), None, "<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>", msg, 0

    def handle_logout():
        return None, gr.update(visible=True), gr.update(visible=False), "*Not Signed In*", gr.update(visible=False), None, None, "<div class='pdf-frame-box empty-preview'>Select or upload a document to preview</div>", [], 0

    auth_outputs = [current_user_state, auth_view, workspace_view, user_badge, btn_logout, doc_selector, active_doc_id, pdf_preview_html, login_status, guest_usage_state]
    btn_do_login.click(handle_login_success, [login_email, login_password], auth_outputs)
    btn_do_register.click(handle_register_success, [reg_email, reg_name, reg_password, reg_confirm_password, reg_role], [current_user_state, auth_view, workspace_view, user_badge, btn_logout, doc_selector, active_doc_id, pdf_preview_html, reg_status, guest_usage_state])
    btn_guest_access.click(handle_guest_access, [], [current_user_state, auth_view, workspace_view, user_badge, btn_logout, doc_selector, active_doc_id, pdf_preview_html, login_status, guest_usage_state])
    btn_logout.click(handle_logout, [], [current_user_state, auth_view, workspace_view, user_badge, btn_logout, doc_selector, active_doc_id, pdf_preview_html, chatbot, guest_usage_state])

    def on_select_document(selected_doc_id, user_data):
        if not selected_doc_id or not user_data:
            return None, "<div class='pdf-frame-box empty-preview'>No document selected</div>"
        return str(selected_doc_id), preview_document_from_db(str(selected_doc_id), user_data)

    doc_selector.change(on_select_document, [doc_selector, current_user_state], [active_doc_id, pdf_preview_html])

    def handle_document_upload(file_obj, u_data, usage_count):
        if not u_data:
            return "Please sign in or continue as guest first.", None, "<div class='pdf-frame-box empty-preview'>Please sign in first.</div>", gr.update(), usage_count
        if not file_obj:
            return "No file selected.", None, "<div class='pdf-frame-box empty-preview'>No document selected</div>", gr.update(), usage_count
        updated_usage, allowed = consume_guest_use(u_data, usage_count)
        if not allowed:
            return guest_limit_message(), None, "<div class='pdf-frame-box error-preview'>Guest limit reached</div>", gr.update(), usage_count
        try:
            msg, doc_id, _, dropdown_update = process_uploaded_document(file_obj, u_data)
            preview = preview_document_from_db(doc_id, u_data)
            return msg, int(doc_id), preview, dropdown_update, updated_usage
        except Exception as exc:
            return f"Error indexing document: {exc}", None, f"<div class='pdf-frame-box error-preview'>{html.escape(str(exc))}</div>", gr.update(), usage_count

    file_input.upload(handle_document_upload, [file_input, current_user_state, guest_usage_state], [status_banner, active_doc_id, pdf_preview_html, doc_selector, guest_usage_state])

    def trigger_quick_chip(prompt_text, chat_hist, doc_id, u_data, usage_count):
        final_result = None
        for result in stream_rag_response(
            prompt_text,
            chat_hist,
            doc_id,
            u_data,
            usage_count,
        ):
            final_result = result

        if final_result is None:
            return chat_hist or [], "", usage_count

        return final_result

    chip_1.click(lambda h, d, u, c: trigger_quick_chip("What are the key takeaways and core ideas from this document?", h, d, u, c), [chatbot, active_doc_id, current_user_state, guest_usage_state], [chatbot, query_input, guest_usage_state])
    chip_2.click(lambda h, d, u, c: trigger_quick_chip("What are the main risks, challenges, and prerequisites mentioned in this document?", h, d, u, c), [chatbot, active_doc_id, current_user_state, guest_usage_state], [chatbot, query_input, guest_usage_state])
    chip_3.click(lambda h, d, u, c: trigger_quick_chip("Summarize the roadmap stages, steps, and structure in this document.", h, d, u, c), [chatbot, active_doc_id, current_user_state, guest_usage_state], [chatbot, query_input, guest_usage_state])
    chip_4.click(lambda h, d, u, c: trigger_quick_chip("What are the recommended action items, next steps, and tools mentioned?", h, d, u, c), [chatbot, active_doc_id, current_user_state, guest_usage_state], [chatbot, query_input, guest_usage_state])

    def handle_chat_submit(user_query, chat_history, doc_id, user_data, usage_count):
        final_result = None
        for result in stream_rag_response(
            user_query,
            chat_history,
            doc_id,
            user_data,
            usage_count,
        ):
            final_result = result

        if final_result is None:
            return chat_history or [], "", usage_count

        return final_result

    btn_send.click(
        fn=handle_chat_submit,
        inputs=[query_input, chatbot, active_doc_id, current_user_state, guest_usage_state],
        outputs=[chatbot, query_input, guest_usage_state],
    )

    query_input.submit(
        fn=handle_chat_submit,
        inputs=[query_input, chatbot, active_doc_id, current_user_state, guest_usage_state],
        outputs=[chatbot, query_input, guest_usage_state],
    )
    btn_clear_chat.click(lambda: ([], ""), [], [chatbot, query_input])

    btn_summarize.click(generate_summary_and_risks, [active_doc_id, current_user_state, guest_usage_state], [out_exec_summary, out_takeaways, out_risks, guest_usage_state])

    def export_summary_to_file(exec_sum, takeaways, risks):
        if not exec_sum or "Click 'Generate" in exec_sum:
            return gr.update(visible=False, value=None)
        path = os.path.join(os.getcwd(), "DocuSense_Summary_Report.md")
        clean_takeaways = re.sub(r"<[^>]+>", "", takeaways or "").replace("&nbsp;", " ")
        clean_risks = re.sub(r"<[^>]+>", "", risks or "").replace("&nbsp;", " ")
        content = f"# DocuSense AI — Document Review Report\n\n## Executive Summary\n{exec_sum}\n\n## Key Takeaways\n{clean_takeaways}\n\n## Risks & Obligations\n{clean_risks}\n"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return gr.update(visible=True, value=path)

    btn_export.click(export_summary_to_file, [out_exec_summary, out_takeaways, out_risks], [export_file])

    def execute_highlight_action(snippet_text: str, action_type: str) -> str:
        prompts = {
            "Explain Simply": f"Explain the following document excerpt in simple terms:\n\n{snippet_text}",
            "Summarize Snippet": f"Summarize the following excerpt:\n\n{snippet_text}",
            "Identify Risks": f"Identify technical, legal, or operational risks in this excerpt:\n\n{snippet_text}",
        }
        return call_llm(prompts[action_type])

    def run_highlight_action(snippet_text, action_type, user_data, usage_count):
        if not (snippet_text or "").strip():
            return "*Please enter or paste a text snippet first.*", usage_count
        updated, allowed = consume_guest_use(user_data, usage_count)
        if not allowed:
            return guest_limit_message(), usage_count
        return execute_highlight_action(snippet_text, action_type), updated

    def call_llm(prompt: str) -> str:
        try:
            result = groq_client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": "You are a professional document review and research assistant."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=1024,
            )
            return result.choices[0].message.content or "No response generated."
        except Exception as exc:
            return f"Error calling LLM: {exc}"

    btn_explain.click(lambda t, u, c: run_highlight_action(t, "Explain Simply", u, c), [snippet_input, current_user_state, guest_usage_state], [snippet_output, guest_usage_state])
    btn_sum_snip.click(lambda t, u, c: run_highlight_action(t, "Summarize Snippet", u, c), [snippet_input, current_user_state, guest_usage_state], [snippet_output, guest_usage_state])
    btn_risk_snip.click(lambda t, u, c: run_highlight_action(t, "Identify Risks", u, c), [snippet_input, current_user_state, guest_usage_state], [snippet_output, guest_usage_state])


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "7860"))

    demo.launch(
        server_name="0.0.0.0",
        server_port=port,
        share=False,
        css=custom_css,
        debug=True,
        show_error=True,
    )
